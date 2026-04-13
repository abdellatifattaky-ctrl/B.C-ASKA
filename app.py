import streamlit as st
import pandas as pd
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm
from io import BytesIO
from datetime import date, timedelta
from num2words import num2words

# --- 1. إعداد قاعدة بيانات الأرشيف ---
ARCHIVE_FILE = "archive_askaoun.csv"

def save_to_archive(num_bc, obj, winner, amount, pv_type):
    if not os.path.isfile(ARCHIVE_FILE):
        df = pd.DataFrame(columns=["Date_Creation", "N_BC", "Objet", "Resultat", "Montant"])
        df.to_csv(ARCHIVE_FILE, index=False, encoding='utf-16')
    
    new_entry = pd.DataFrame([[date.today().strftime('%d/%m/%Y'), num_bc, obj, winner, amount]], 
                             columns=["Date_Creation", "N_BC", "Objet", "Resultat", "Montant"])
    new_entry.to_csv(ARCHIVE_FILE, mode='a', header=False, index=False, encoding='utf-16')

# --- 2. دالة تحويل المبالغ إلى حروف ---
def format_to_words_fr(amount_str):
    try:
        val = float(str(amount_str).replace(' ', '').replace(',', ''))
        integer_part = int(val)
        cents = int(round((val - integer_part) * 100))
        words = num2words(integer_part, lang='fr').upper()
        text = f"{words} DHS TTC"
        if cents > 0:
            cents_words = num2words(cents, lang='fr').upper()
            text += f" ,{cents_words} CENTIMES"
        else:
            text += " ,00CTS"
        return text
    except:
        return "________________"

# --- 3. إعدادات الصفحة والذاكرة (Session State) ---
st.set_page_config(page_title="Commune Askaouen - Système PV", layout="wide")

if 'logo_data' not in st.session_state:
    st.session_state['logo_data'] = None

# --- 4. الشريط الجانبي (الشعار + اللجنة) ---
st.sidebar.header("🖼️ الهوية البصرية")
uploaded_logo = st.sidebar.file_uploader("ارفع شعار الجماعة (يبقى محفوظاً)", type=['png', 'jpg', 'jpeg'])
if uploaded_logo:
    st.session_state['logo_data'] = uploaded_logo.getvalue()

st.sidebar.divider()
st.sidebar.header("👥 اختيار أعضاء اللجنة")
all_members = [
    {"id": "p", "name": "MOHAMED ZILALI", "role": "Président de la commission"},
    {"id": "d", "name": "M BAREK BAK", "role": "Directeur du service"},
    {"id": "t1", "name": "ATTAKY ABDELLATIF", "role": "Technicien de la commune"},
    {"id": "s", "name": "NOUREDDIN SALHI", "role": "Service des dépenses"},
    {"id": "t2", "name": "FAYSSAL KADRI", "role": "Technicien à la commune"}
]

selected_members = []
for member in all_members:
    if st.sidebar.checkbox(member['name'], value=True, key=member['id']):
        selected_members.append(member)

# --- 5. نظام التبويبات (Tabs) ---
tab_main, tab_archive = st.tabs(["📝 إعداد المحاضر", "📂 الأرشيف التاريخي"])

with tab_main:
    st.title("🏛️ نظام استخراج المحاضر - جماعة أسكاون")
    
    with st.expander("📝 Détails Administratifs", expanded=True):
        c1, c2 = st.columns(2)
        num_bc = c1.text_input("N° BC", "01/ASK/2025")
        date_pub = c2.date_input("Date de publication", date(2025, 3, 25))
        obj_bc = st.text_area("Objet", "Location d’une Tractopelle pour les travaux divers.")

    st.subheader("📊 Liste des concurrents")
    df_init = pd.DataFrame([
        {"Rang": 1, "Nom": "STE OUBRAIM SARL", "Montant": "69840.00"},
        {"Rang": 2, "Nom": "DECO GRC", "Montant": "93120.00"},
        {"Rang": 3, "Nom": "AIT MOUMOU REALISATION ET CONSTRUCTION", "Montant": "102432.00"},
        {"Rang": 4, "Nom": "KADEM SARL", "Montant": "111744.00"},
        {"Rang": 5, "Nom": "TOUZANI 2ZD", "Montant": "114072.00"}
    ])
    data = st.data_editor(df_init, use_container_width=True)

    st.divider()

    c_v1, c_v2, c_v3 = st.columns(3)
    pv_num = c_v1.selectbox("Numéro du PV:", [1, 2, 3, 4, 5, 6])
    reunion_date = c_v3.date_input("Date de la séance", date.today())
    reunion_hour = c_v2.text_input("Heure", "12h00mn")

    # --- ميزة الموعد القادم التلقائي (J+1) ---
    next_meeting_default = reunion_date + timedelta(days=1)
    next_meeting = st.date_input("Date du prochain rendez-vous (J+1)", next_meeting_default)

    is_final_attr = False
    is_infructueux = False
    if pv_num == 6:
        res_6 = st.radio("Résultat final:", ["Attribution (إسناد)", "B.C Infructueux (بدون جدوى)"])
        is_infructueux = (res_6 == "B.C Infructueux (بدون جدوى)")
        is_final_attr = (res_6 == "Attribution (إسناد)")
    else:
        is_final_attr = st.checkbox("✅ Est-ce le PV d'attribution finale ? (إسناد)")

    if st.button("🚀 إنشاء المحضر وحفظ البيانات"):
        if not selected_members:
            st.error("المرجو اختيار أعضاء اللجنة من القائمة الجانبية!")
        else:
            doc = Document()
            section = doc.sections[0]
            section.page_height, section.page_width = Cm(29.7), Cm(21)
            section.left_margin, section.right_margin = Cm(2.5), Cm(2)
            section.top_margin, section.bottom_margin = Cm(1.5), Cm(1.5)

            # --- الترويسة الاحترافية مع الشعار ---
            header = section.header
            htable = header.add_table(1, 3, Inches(6.5))
            # يمين (عربي)
            c_right = htable.rows[0].cells[2].paragraphs[0]
            c_right.text = "المملكة المغربية\nوزارة الداخلية\nجماعة أسكاون"
            c_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            c_right.style.font.size = Pt(9)
            # وسط (شعار)
            if st.session_state['logo_data']:
                logo_stream = BytesIO(st.session_state['logo_data'])
                logo_run = htable.rows[0].cells[1].paragraphs[0].add_run()
                logo_run.add_picture(logo_stream, width=Cm(1.9))
                htable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # يسار (فرنسي)
            c_left = htable.rows[0].cells[0].paragraphs[0]
            c_left.text = "ROYAUME DU MAROC\nMINISTERE DE L'INTERIEUR\nCOMMUNE D'ASKAOUN"
            c_left.style.font.size = Pt(9)

            # --- نص المحضر ---
            pv_label = "1ére" if pv_num == 1 else f"{pv_num}éme"
            title = doc.add_heading(f"{pv_label} Procès verbal", 1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_sub = doc.add_paragraph("De la commission d’ouverture des plis\nProcédure Bon de commande")
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.runs[0].bold = True

            doc.add_paragraph(f"Objet : {obj_bc}").bold = True
            doc.add_paragraph(f"Le {reunion_date.strftime('%d/%m/%Y')} à {reunion_hour}, la commission composée de :")
            for m in selected_members:
                doc.add_paragraph(f"- {m['name']} : {m['role']}")

            doc.add_paragraph(f"\nS’est réunie في salle de réunion de la commune بشأن avis d’achat du bon de commande n° {num_bc} publié le : {date_pub.strftime('%d/%m/%Y')} على بوابة الصفقات العمومية، ayant pour objet : {obj_bc}")

            # منطق اختيار الشركات
            idx = min(pv_num - 1, len(data) - 1)
            curr_co = data.iloc[idx]
            amt_words = format_to_words_fr(curr_co['Montant'])

            if pv_num == 1:
                doc.add_paragraph("Les soumissionnaires qui ont déposés leurs offres électroniquement sont :")
                tab = doc.add_table(rows=1, cols=3); tab.style = 'Table Grid'
                for _, r in data.iterrows():
                    row = tab.add_row().cells
                    row[0].text, row[1].text, row[2].text = str(r['Rang']), r['Nom'], f"{r['Montant']} MAD"
                doc.add_paragraph("\nFormat papier : Néant.")
                doc.add_paragraph(f"Le président invite {curr_co['Nom']} ({amt_words}) لتأكيد عرضه، ورفع الجلسة إلى غاية {next_meeting.strftime('%d/%m/%Y')}.")

            elif is_infructueux:
                doc.add_paragraph(f"تبين أن شركة {curr_co['Nom']} لم تؤكد عرضها.")
                p_inf = doc.add_paragraph("\nPAR CONSEQUENT, LA COMMISSION DECLARE QUE CE BON DE COMMANDE EST :")
                p_inf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                res_inf = doc.add_paragraph("INFRUCTUEUX")
                res_inf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                res_inf.bold = True

            elif is_final_attr:
                doc.add_paragraph(f"تبين أن شركة : {curr_co['Nom']} قد أكدت عرضها.")
                p_res = doc.add_paragraph(f"Le président VALIDE la confirmation et ATTRIBUE le BC à {curr_co['Nom']} بمبلغ {curr_co['Montant']} Dhs TTC {amt_words}.")
                p_res.bold = True

            else:
                prev_co = data.iloc[idx - 1]
                doc.add_paragraph(f"تبين أن شركة {prev_co['Nom']} لم تؤكد عرضها، وبالتالي تم استبعادها.")
                doc.add_paragraph(f"تمت دعوة الشركة التالية: {curr_co['Nom']} (رتبة {pv_num}) بمبلغ {curr_co['Montant']} Dhs ({amt_words}) لتأكيد العرض في موعد {next_meeting.strftime('%d/%m/%Y')}.")

            # --- التوقيعات ---
            doc.add_paragraph(f"\nAskaouen le {reunion_date.strftime('%d/%m/%Y')}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig_tab = doc.add_table(rows=(len(selected_members)//2)+1, cols=2)
            for i, m in enumerate(selected_members):
                cell = sig_tab.cell(i//2, i%2)
                cell.text = m['name']
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # حفظ الأرشيف
            res_val = curr_co['Nom'] if is_final_attr else ("Infructueux" if is_infructueux else "En cours")
            save_to_archive(num_bc, obj_bc, res_val, curr_co['Montant'], pv_num)

            bio = BytesIO()
            doc.save(bio)
            st.success("✅ تم توليد المحضر بنجاح!")
            st.download_button(f"📥 تحميل PV {pv_num}", bio.getvalue(), f"PV_{num_bc.replace('/','_')}.docx")

with tab_archive:
    st.header("📂 أرشيف سندات الطلب")
    if os.path.isfile(ARCHIVE_FILE):
        archive_df = pd.read_csv(ARCHIVE_FILE, encoding='utf-16')
        st.dataframe(archive_df, use_container_width=True)
        st.download_button("📥 تحميل السجل الكامل (CSV)", archive_df.to_csv(index=False).encode('utf-16'), "Archive_Askaouen.csv")
    else:
        st.info("لا يوجد بيانات في الأرشيف حالياً.")
